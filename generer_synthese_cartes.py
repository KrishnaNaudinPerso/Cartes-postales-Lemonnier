#!/usr/bin/env python3
"""
=============================================================================
SYNTHÈSE DES CARTES POSTALES LEMONNIER  —  v2 (exécution par morceaux)
=============================================================================

Le script s'exécute en étapes indépendantes, une carte à la fois.
Chaque carte est sauvegardée immédiatement après traitement.
En cas de timeout ou d'erreur, relancez simplement la même commande.

PRÉREQUIS :
  pip install anthropic python-docx Pillow

CLÉ API :
  export ANTHROPIC_API_KEY="sk-ant-..."

UTILISATION :
  # Étape 1 – Analyser les cartes (une par une, relançable)
  python3 generer_synthese_cartes.py analyser

  # Analyser seulement certaines cartes
  python3 generer_synthese_cartes.py analyser --cartes 1,5,100

  # Analyser par tranche (ex : cartes 1 à 50)
  python3 generer_synthese_cartes.py analyser --de 1 --a 50

  # Étape 2 – Générer le document Word (sans appel API)
  python3 generer_synthese_cartes.py generer

  # Voir l'état d'avancement
  python3 generer_synthese_cartes.py statut

  # Réanalyser une carte spécifique (même si déjà traitée)
  python3 generer_synthese_cartes.py analyser --cartes 42 --forcer
=============================================================================
"""

import os
import re
import sys
import json
import base64
import argparse
import time
from pathlib import Path
from collections import defaultdict

# ─── CONFIGURATION ────────────────────────────────────────────────────────────

DOSSIER = Path(__file__).parent
API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
FICHIER_SAUVEGARDE = DOSSIER / "progression_analyse.json"
FICHIER_SORTIE     = DOSSIER / "synthese_cartes_lemonnier.docx"
MODELE_CLAUDE      = "claude-opus-4-5-20251101"

# Taille max des images envoyées (pixels) — réduire si timeouts fréquents
MAX_PIXELS = 1200

# Pause entre appels API (secondes)
PAUSE_API = 1.5

# Nombre de tentatives en cas d'erreur/timeout
MAX_TENTATIVES = 3

# ─── IMPORTS ─────────────────────────────────────────────────────────────────

try:
    import anthropic
except ImportError:
    print("❌ Installer : pip install anthropic")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ Installer : pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    import io as _io
    PILLOW_OK = True
except ImportError:
    PILLOW_OK = False


# =============================================================================
# UTILITAIRES
# =============================================================================

def charger_sauvegarde() -> dict:
    if FICHIER_SAUVEGARDE.exists():
        with open(FICHIER_SAUVEGARDE, encoding="utf-8") as f:
            return json.load(f)
    return {}


def sauvegarder(analyses: dict):
    with open(FICHIER_SAUVEGARDE, "w", encoding="utf-8") as f:
        json.dump(analyses, f, ensure_ascii=False, indent=2, default=str)


def encoder_image(chemin: Path) -> tuple:
    """Encode l'image en base64, réduite si nécessaire."""
    with open(chemin, "rb") as f:
        donnees = f.read()

    if PILLOW_OK:
        img = Image.open(_io.BytesIO(donnees))
        if max(img.size) > MAX_PIXELS:
            img.thumbnail((MAX_PIXELS, MAX_PIXELS), Image.LANCZOS)
            buf = _io.BytesIO()
            img.save(buf, format="JPEG", quality=82)
            donnees = buf.getvalue()

    b64 = base64.standard_b64encode(donnees).decode("utf-8")
    ext = chemin.suffix.lower()
    mime = {".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".png": "image/png", ".tif": "image/tiff",
            ".tiff": "image/tiff"}.get(ext, "image/jpeg")
    return b64, mime


def appel_api(client, messages, max_tokens=900, tentative=1) -> str:
    """Appel API avec retry automatique."""
    try:
        rep = client.messages.create(
            model=MODELE_CLAUDE,
            max_tokens=max_tokens,
            messages=messages
        )
        return rep.content[0].text.strip()
    except Exception as e:
        msg = str(e)
        if tentative < MAX_TENTATIVES:
            attente = tentative * 5
            print(f"    ⚠️  Erreur ({msg[:60]}...) → Réessai dans {attente}s")
            time.sleep(attente)
            return appel_api(client, messages, max_tokens, tentative + 1)
        print(f"    ❌ Échec après {MAX_TENTATIVES} tentatives : {msg[:80]}")
        return ""


def parse_json_reponse(texte: str) -> dict:
    """Extrait le JSON d'une réponse, même entourée de markdown."""
    if not texte:
        return {}
    texte = re.sub(r"```json\s*", "", texte)
    texte = re.sub(r"```\s*", "", texte)
    # Trouver le premier { et le dernier }
    debut = texte.find("{")
    fin   = texte.rfind("}")
    if debut == -1 or fin == -1:
        return {"texte_brut": texte}
    try:
        return json.loads(texte[debut:fin+1])
    except json.JSONDecodeError:
        return {"texte_brut": texte[debut:fin+1]}


# =============================================================================
# PHASE 1 : GEDCOM
# =============================================================================

def lire_gedcom(chemin_ged: Path) -> dict:
    print("\n📖 Lecture du fichier GEDCOM...")
    personnes = {}
    familles  = {}

    if not chemin_ged.exists():
        print("  ⚠️  Fichier GED non trouvé")
        return personnes

    def to_iso(s):
        mois = {"JAN":"01","FEB":"02","MAR":"03","APR":"04","MAY":"05",
                "JUN":"06","JUL":"07","AUG":"08","SEP":"09","OCT":"10",
                "NOV":"11","DEC":"12"}
        s = re.sub(r"(?i)(ABT|BEF|AFT|CAL|EST|CIRCA|ABOUT|~)\s*", "", s).strip()
        p = s.split()
        try:
            if len(p) == 3:
                return f"{p[2]}-{mois.get(p[1].upper(),'00')}-{int(p[0]):02d}"
            if len(p) == 2:
                return f"{p[1]}-{mois.get(p[0].upper(),'00')}-00"
            if len(p) == 1:
                return f"{p[0]}-00-00"
        except (ValueError, IndexError):
            pass
        return s

    with open(chemin_ged, encoding="utf-8", errors="replace") as f:
        lignes = f.readlines()

    ind = None
    fam = None
    tag1 = None

    for ligne in lignes:
        ligne = ligne.rstrip()
        parts = ligne.split(" ", 2)
        if len(parts) < 2:
            continue
        niv = parts[0].strip()
        tag = parts[1].strip()
        val = parts[2].strip() if len(parts) > 2 else ""

        if niv == "0":
            ind = fam = tag1 = None
            if "@" in tag and val == "INDI":
                ind = tag
                personnes[ind] = {"id": ind, "nom_complet": "", "prenom": "",
                                  "nom": "", "sexe": "", "naissance": "",
                                  "deces": "", "lieu_naissance": "",
                                  "lieu_deces": "", "famille_enfant": [],
                                  "famille_conjoint": []}
            elif "@" in tag and val == "FAM":
                fam = tag
                familles[fam] = {"mari": "", "femme": "", "enfants": []}

        elif niv == "1":
            tag1 = tag
            if ind:
                p = personnes[ind]
                if tag == "NAME":
                    pts = val.split("/")
                    p["prenom"] = pts[0].strip() if pts else ""
                    p["nom"]    = pts[1].strip() if len(pts) > 1 else ""
                    p["nom_complet"] = val.replace("/", " ").strip()
                elif tag == "SEX":
                    p["sexe"] = val
                elif tag == "FAMC":
                    p["famille_enfant"].append(val)
                elif tag == "FAMS":
                    p["famille_conjoint"].append(val)
            elif fam:
                f = familles[fam]
                if tag == "HUSB": f["mari"] = val
                elif tag == "WIFE": f["femme"] = val
                elif tag == "CHIL": f["enfants"].append(val)

        elif niv == "2" and ind and tag1:
            p = personnes[ind]
            if tag1 == "BIRT":
                if tag == "DATE": p["naissance"] = to_iso(val)
                elif tag == "PLAC": p["lieu_naissance"] = val.split(",")[0].strip()
            elif tag1 == "DEAT":
                if tag == "DATE": p["deces"] = to_iso(val)
                elif tag == "PLAC": p["lieu_deces"] = val.split(",")[0].strip()

    # Relations
    for fid, f in familles.items():
        for enfant_id in f["enfants"]:
            if enfant_id in personnes:
                if f["mari"]  in personnes: personnes[enfant_id]["_pere"] = f["mari"]
                if f["femme"] in personnes: personnes[enfant_id]["_mere"] = f["femme"]
        if f["mari"] in personnes and f["femme"] in personnes:
            personnes[f["mari"]]["_conjoint"]  = f["femme"]
            personnes[f["femme"]]["_conjoint"] = f["mari"]

    print(f"  ✅ {len(personnes)} personnes, {len(familles)} familles")
    return personnes


# =============================================================================
# PHASE 2 : GROUPEMENT DES IMAGES
# =============================================================================

def grouper_images(dossier: Path) -> dict:
    print("\n🗂️  Groupement des images...")

    cartes = defaultdict(lambda: {"recto": None, "verso": None,
                                   "type": "carte", "autres": []})

    patterns = [
        (r"(?i)carte\s*(\d+)\s*\((\d+)\)",         "carte"),
        (r"(?i)graphe\s+(\d+)\s*[\s.(]+(\d+)",      "graphe"),
        (r"(?i)carte\s*(\d+)\.(\d+)",               "carte"),
        (r"(?i)carte\s+(\d+)\s+(\d+)",              "carte"),
        (r"^(\d+)\s+(\d+)$",                        "carte"),
        (r"(?i)carte\s*(\d+)$",                     "carte_sans_face"),
    ]

    non_matches = []
    extensions_img = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ""}

    for fichier in sorted(dossier.iterdir()):
        if not fichier.is_file():
            continue
        if fichier.suffix.lower() not in extensions_img:
            continue
        if fichier.suffix.lower() in {".docx", ".ged", ".json", ".py"}:
            continue

        nom = fichier.stem
        matched = False

        for pattern, type_c in patterns:
            m = re.search(pattern, nom)
            if m:
                num = int(m.group(1))
                if type_c == "carte_sans_face":
                    cartes[num]["type"] = "carte"
                    if cartes[num]["recto"] is None:
                        cartes[num]["recto"] = fichier
                    else:
                        cartes[num]["autres"].append(fichier)
                    matched = True
                    break

                face = int(m.group(2))
                cartes[num]["type"] = type_c
                if face == 1:
                    if cartes[num]["recto"] is None:
                        cartes[num]["recto"] = fichier
                    else:
                        cartes[num]["autres"].append(fichier)
                elif face == 2:
                    if cartes[num]["verso"] is None:
                        cartes[num]["verso"] = fichier
                    else:
                        cartes[num]["autres"].append(fichier)
                matched = True
                break

        if not matched:
            non_matches.append(fichier.name)

    total = len(cartes)
    complets = sum(1 for c in cartes.values() if c["recto"] and c["verso"])
    print(f"  ✅ {total} cartes ({complets} complètes recto+verso)")
    if non_matches:
        print(f"  ⚠️  {len(non_matches)} fichier(s) non identifié(s) : {non_matches[:5]}")

    return dict(sorted(cartes.items()))


# =============================================================================
# PHASE 3 : ANALYSE — RECTO (appel séparé, léger)
# =============================================================================

PROMPT_RECTO = """Tu analyses le RECTO d'une carte postale ancienne (famille Lemonnier, Mayenne, France).
Réponds UNIQUEMENT avec ce JSON compact (pas de markdown autour) :
{
  "lieu": "ville/lieu représenté",
  "description": "description concise de l'image (max 2 phrases)",
  "editeur": "éditeur imprimeur si lisible, sinon vide",
  "periode_estimee": "décennie estimée ex: 1910s, 1930s"
}"""

PROMPT_VERSO = """Tu analyses le VERSO d'une carte postale ancienne française.
La famille Lemonnier est originaire de Craon/Laval (Mayenne). Guy Lemonnier est né en 1903.

Réponds UNIQUEMENT avec ce JSON compact (pas de markdown autour) :
{
  "transcription": "transcription FIDÈLE et COMPLÈTE du manuscrit. Mettre [illisible] si un mot est impossible à lire. Ne PAS inventer de mots.",
  "expediteur": "prénom et nom de l'expéditeur si lisible",
  "destinataire": "prénom et nom du destinataire si lisible",
  "adresse": "adresse postale complète si lisible",
  "date_manuscrite": "date écrite à la main si présente",
  "cachet_date": "date lue sur le cachet postal (JJ/MM/AAAA ou MM/AAAA)",
  "cachet_ville": "ville d'expédition du cachet",
  "timbre": "couleur, valeur, thème du timbre"
}"""

PROMPT_CONTEXTE = """Une carte postale ancienne a été envoyée avec ces informations :
- Date cachet : {date}
- Ville expédition : {ville}
- Timbre : {timbre}
- Lieu représenté : {lieu}

La famille centrale : Guy Lemonnier (né 1903 à Ernée, Mayenne, décédé 1979),
sa femme Augusta Boulay (née 1907 à Craon).

Réponds UNIQUEMENT avec ce JSON compact (pas de markdown) :
{
  "contexte_historique": "1-2 événements majeurs en France/Europe à cette date",
  "vie_quotidienne": "1 phrase sur la vie quotidienne en France à cette époque",
  "philatelie": "identification du timbre (série, date émission) si reconnaissable",
  "ville_info": "informations sur la ville d'expédition à l'époque (département, particularités)"
}"""


def analyser_recto(client, chemin_img: Path) -> dict:
    """Analyse séparée du recto — appel léger."""
    b64, mime = encoder_image(chemin_img)
    texte = appel_api(client, [{
        "role": "user",
        "content": [
            {"type": "image", "source": {"type": "base64",
                                          "media_type": mime, "data": b64}},
            {"type": "text",  "text": PROMPT_RECTO}
        ]
    }], max_tokens=400)
    time.sleep(PAUSE_API)
    return parse_json_reponse(texte)


def analyser_verso(client, chemin_img: Path) -> dict:
    """Analyse séparée du verso — appel principal."""
    b64, mime = encoder_image(chemin_img)
    texte = appel_api(client, [{
        "role": "user",
        "content": [
            {"type": "image", "source": {"type": "base64",
                                          "media_type": mime, "data": b64}},
            {"type": "text",  "text": PROMPT_VERSO}
        ]
    }], max_tokens=700)
    time.sleep(PAUSE_API)
    return parse_json_reponse(texte)


def enrichir_contexte(client, recto: dict, verso: dict) -> dict:
    """Enrichissement historique et philatélique — texte pur, pas d'image."""
    prompt = PROMPT_CONTEXTE.format(
        date=verso.get("cachet_date", "inconnue"),
        ville=verso.get("cachet_ville", "inconnue"),
        timbre=verso.get("timbre", "inconnu"),
        lieu=recto.get("lieu", "inconnu")
    )
    texte = appel_api(client, [{"role": "user", "content": prompt}],
                      max_tokens=500)
    time.sleep(PAUSE_API)
    return parse_json_reponse(texte)


def calculer_ages(personnes: dict, verso: dict) -> dict:
    """Calcule les âges des membres clés à la date de la carte."""
    ages = {}
    date_str = verso.get("cachet_date", "") or verso.get("date_manuscrite", "")
    if not date_str:
        return ages

    m = re.search(r"\b(19\d{2}|20\d{2})\b", date_str)
    if not m:
        return ages
    annee = int(m.group(1))

    noms_cles = ["Lemonnier", "Boulay", "Naudin", "Crombach"]
    for pid, p in personnes.items():
        if not any(n in p.get("nom", "") for n in noms_cles):
            continue
        naissance = p.get("naissance", "")
        if naissance and len(naissance) >= 4 and naissance[:4].isdigit():
            age = annee - int(naissance[:4])
            if 0 < age < 110:
                ages[p["nom_complet"]] = age

    return ages


# =============================================================================
# PHASE 4 : ANALYSE D'UNE CARTE (orchestration)
# =============================================================================

def analyser_carte(client, num: int, info: dict, personnes: dict) -> dict:
    """
    Analyse complète d'une carte en 3 petits appels API distincts.
    Retourne un dict avec toutes les données.
    """
    recto_data  = {}
    verso_data  = {}
    contexte    = {}

    # — Appel 1 : Recto ───────────────────────────────────────────────────────
    if info["recto"]:
        print(f"    📷 Recto...", end=" ", flush=True)
        recto_data = analyser_recto(client, info["recto"])
        print("✓")
    else:
        print(f"    📷 Recto absent")

    # — Appel 2 : Verso ───────────────────────────────────────────────────────
    if info["verso"]:
        print(f"    ✍️  Verso...", end=" ", flush=True)
        verso_data = analyser_verso(client, info["verso"])
        print("✓")
    else:
        print(f"    ✍️  Verso absent")

    # — Appel 3 : Contexte (texte seulement, rapide) ──────────────────────────
    if recto_data or verso_data:
        print(f"    🏛  Contexte...", end=" ", flush=True)
        contexte = enrichir_contexte(client, recto_data, verso_data)
        print("✓")

    # — Calcul des âges ───────────────────────────────────────────────────────
    ages = calculer_ages(personnes, verso_data)

    return {
        "num": num,
        "recto":    recto_data,
        "verso":    verso_data,
        "contexte": contexte,
        "ages":     ages,
        "fichiers": {
            "recto": str(info["recto"]) if info["recto"] else None,
            "verso": str(info["verso"]) if info["verso"] else None,
        }
    }


# =============================================================================
# COMMANDE : analyser
# =============================================================================

def cmd_analyser(args):
    api_key = API_KEY or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("❌ ANTHROPIC_API_KEY manquante (export ANTHROPIC_API_KEY='sk-ant-...')")
        sys.exit(1)

    client     = anthropic.Anthropic(api_key=api_key)
    personnes  = lire_gedcom(next(DOSSIER.glob("*.ged"), Path("inexistant.ged")))
    cartes     = grouper_images(DOSSIER)
    sauvegarde = charger_sauvegarde()

    # Déterminer quelles cartes traiter
    if args.cartes:
        nums = [int(n.strip()) for n in args.cartes.split(",")]
    else:
        debut = args.de if args.de else min(cartes.keys())
        fin   = args.a  if args.a  else max(cartes.keys())
        nums  = [n for n in sorted(cartes.keys()) if debut <= n <= fin]

    # Filtrer celles déjà traitées (sauf --forcer)
    if not args.forcer:
        nums_restants = [n for n in nums if str(n) not in sauvegarde]
    else:
        nums_restants = nums

    if not nums_restants:
        print("\n✅ Toutes les cartes demandées sont déjà analysées.")
        print("   Utilisez --forcer pour réanalyser.")
        return

    print(f"\n🔍 {len(nums_restants)} carte(s) à analyser "
          f"({len(sauvegarde)} déjà faites sur {len(cartes)} total)")
    print("   Sauvegarde après chaque carte. Relançable en cas d'interruption.\n")

    for i, num in enumerate(nums_restants, 1):
        if num not in cartes:
            print(f"  ⚠️  Carte {num} non trouvée dans les fichiers")
            continue

        print(f"  [{i}/{len(nums_restants)}] Carte n°{num}")
        try:
            resultat = analyser_carte(client, num, cartes[num], personnes)
            sauvegarde[str(num)] = resultat
            sauvegarder(sauvegarde)
            print(f"    💾 Sauvegardé\n")
        except KeyboardInterrupt:
            print("\n  ⛔ Interruption. Progression sauvegardée.")
            break
        except Exception as e:
            print(f"    ❌ Erreur carte {num} : {e}")
            sauvegarde[str(num)] = {"erreur": str(e), "num": num,
                                     "fichiers": {
                                         "recto": str(cartes[num]["recto"]) if cartes[num]["recto"] else None,
                                         "verso": str(cartes[num]["verso"]) if cartes[num]["verso"] else None,
                                     }}
            sauvegarder(sauvegarde)
            print(f"    (Erreur sauvegardée, on continue)\n")
            time.sleep(2)

    total_fait = len(sauvegarde)
    print(f"\n✅ Analyse terminée — {total_fait}/{len(cartes)} cartes traitées")
    if total_fait < len(cartes):
        print("   Pour continuer : python3 generer_synthese_cartes.py analyser")
    else:
        print("   Pour générer le document : python3 generer_synthese_cartes.py generer")


# =============================================================================
# COMMANDE : generer (document Word)
# =============================================================================

def creer_style(styles, nom, base, taille, gras=False, couleur=None, av=0, ap=0):
    try:
        style = styles.add_style(nom, 1)
    except ValueError:
        style = styles[nom]
    style.base_style = styles.get(base, styles["Normal"])
    style.font.size  = Pt(taille)
    style.font.bold  = gras
    if couleur:
        style.font.color.rgb = RGBColor(*couleur)
    style.paragraph_format.space_before = Pt(av)
    style.paragraph_format.space_after  = Pt(ap)
    return style


def cmd_generer(args):
    sauvegarde = charger_sauvegarde()
    if not sauvegarde:
        print("❌ Aucune analyse trouvée. Lancez d'abord : analyser")
        return

    personnes = lire_gedcom(next(DOSSIER.glob("*.ged"), Path("inexistant.ged")))

    print(f"\n📄 Génération du document ({len(sauvegarde)} cartes)...")

    doc = Document()

    # Mise en page A4
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2)
    sec.top_margin  = sec.bottom_margin = Cm(2)

    # Styles
    st = doc.styles
    creer_style(st, "TitreCarte", "Normal", 15, gras=True,
                couleur=(80, 50, 10), av=8, ap=4)
    creer_style(st, "SousTitre2", "Normal", 11, gras=True,
                couleur=(110, 70, 20), av=6, ap=2)
    creer_style(st, "Transcr",    "Normal", 10.5, couleur=(20, 20, 20), av=2, ap=2)
    creer_style(st, "ContexteH",  "Normal", 10,   couleur=(40, 80, 40), av=2, ap=2)
    creer_style(st, "InfoGrise",  "Normal",  9.5, couleur=(90, 90, 90), av=1, ap=1)

    # ── Page de titre ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cartes Postales — Famille Lemonnier")
    r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(80, 50, 10)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Archives familiales · Document de synthèse")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(120, 100, 70)

    doc.add_paragraph()

    # Intro généalogique
    guy_id = next((pid for pid, p in personnes.items()
                   if "Guy" in p.get("prenom","") and "Lemonnier" in p.get("nom","")), None)
    if guy_id:
        g = personnes[guy_id]
        nb = g["naissance"][:4] if g["naissance"] else "?"
        nd = g["deces"][:4]    if g["deces"]     else "?"
        p3 = doc.add_paragraph(
            f"Guy Michel Albert Lemonnier ({nb}–{nd}), né à {g['lieu_naissance']}, "
            f"est le personnage central de cette correspondance familiale. "
            f"Ce document présente {len(sauvegarde)} cartes postales analysées, "
            f"avec transcription du manuscrit, identification des cachets "
            f"et éclairage historique."
        )
        p3.style = st["InfoGrise"]

    doc.add_paragraph()
    pi = doc.add_paragraph(f"Généré le {__import__('datetime').date.today().strftime('%d/%m/%Y')}")
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.style = st["InfoGrise"]

    # ── Une page par carte ────────────────────────────────────────────────────
    for num_str in sorted(sauvegarde.keys(), key=lambda x: int(x)):
        d = sauvegarde[num_str]
        num = int(num_str)

        recto   = d.get("recto", {})
        verso   = d.get("verso", {})
        ctx     = d.get("contexte", {})
        ages    = d.get("ages", {})
        fichiers= d.get("fichiers", {})
        erreur  = d.get("erreur")

        doc.add_page_break()

        # Titre de la carte
        ville  = verso.get("cachet_ville") or recto.get("lieu") or "?"
        date_c = verso.get("cachet_date")  or recto.get("periode_estimee") or "?"
        exped  = verso.get("expediteur", "?")
        dest   = verso.get("destinataire", "?")

        p_titre = doc.add_paragraph(style=st["TitreCarte"])
        p_titre.add_run(f"Carte n°{num}  —  {ville}  —  {date_c}")

        if exped != "?" or dest != "?":
            p_route = doc.add_paragraph(style=st["SousTitre2"])
            p_route.add_run(f"De : {exped}     →     À : {dest}")

        if erreur:
            doc.add_paragraph(f"⚠️ Erreur lors de l'analyse : {erreur}").style = st["InfoGrise"]

        # Images côte à côte
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        def ajouter_img(cell, chemin_str, legende):
            p_l = cell.add_paragraph(legende)
            p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_l = p_l.runs[0]
            run_l.font.bold = True
            run_l.font.size = Pt(9)
            if chemin_str and Path(chemin_str).exists():
                p_i = cell.add_paragraph()
                p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p_i.add_run().add_picture(chemin_str, width=Cm(8.5))
                except Exception as e:
                    p_i.add_run(f"[Image non chargée : {e}]")

        ajouter_img(table.cell(0, 0), fichiers.get("recto"), "RECTO")
        ajouter_img(table.cell(0, 1), fichiers.get("verso"),  "VERSO")

        doc.add_paragraph()

        # Description recto
        if recto.get("description"):
            doc.add_paragraph("🏞  Recto", style=st["SousTitre2"])
            doc.add_paragraph(recto["description"], style=st["InfoGrise"])
            if recto.get("editeur"):
                doc.add_paragraph(f"Éditeur : {recto['editeur']}", style=st["InfoGrise"])

        # Transcription
        transcription = verso.get("transcription", "")
        if transcription:
            doc.add_paragraph("✍  Transcription du manuscrit", style=st["SousTitre2"])

            p_t = doc.add_paragraph(style=st["Transcr"])
            # Bordure légère autour de la transcription
            pPr = p_t._p.get_or_add_pPr()
            pPrBdr = OxmlElement("w:pBdr")
            for bord in ["top", "left", "bottom", "right"]:
                b = OxmlElement(f"w:{bord}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "4")
                b.set(qn("w:color"), "C8A45A")
                pPrBdr.append(b)
            pPr.append(pPrBdr)

            r_t = p_t.add_run(transcription)
            r_t.font.name   = "Georgia"
            r_t.font.size   = Pt(10.5)
            r_t.font.italic = True

        # Informations postales
        doc.add_paragraph("📮  Cachet et timbre", style=st["SousTitre2"])
        infos = []
        if verso.get("cachet_date"):
            infos.append(f"Date du cachet : {verso['cachet_date']}")
        if verso.get("cachet_ville"):
            infos.append(f"Ville d'expédition : {verso['cachet_ville']}")
        if verso.get("date_manuscrite"):
            infos.append(f"Date manuscrite : {verso['date_manuscrite']}")
        if verso.get("timbre"):
            infos.append(f"Timbre : {verso['timbre']}")
        if ctx.get("philatelie"):
            infos.append(f"Identification philatélique : {ctx['philatelie']}")
        if verso.get("adresse"):
            infos.append(f"Adresse destinataire : {verso['adresse']}")
        for info in infos:
            doc.add_paragraph(f"• {info}", style=st["InfoGrise"])
        if not infos:
            doc.add_paragraph("(Cachet illisible ou absent)", style=st["InfoGrise"])

        # Contexte historique
        if ctx.get("contexte_historique") or ctx.get("vie_quotidienne"):
            doc.add_paragraph("🏛  Contexte historique", style=st["SousTitre2"])
            if ctx.get("contexte_historique"):
                doc.add_paragraph(ctx["contexte_historique"], style=st["ContexteH"])
            if ctx.get("vie_quotidienne"):
                doc.add_paragraph(ctx["vie_quotidienne"], style=st["ContexteH"])

        # Ville d'expédition
        if ctx.get("ville_info"):
            doc.add_paragraph(f"📍  {verso.get('cachet_ville','Lieu d envoi')}",
                              style=st["SousTitre2"])
            doc.add_paragraph(ctx["ville_info"], style=st["InfoGrise"])

        # Âges
        if ages:
            doc.add_paragraph("👪  Âges à l'envoi", style=st["SousTitre2"])
            for nom, age in ages.items():
                doc.add_paragraph(f"• {nom} : {age} ans", style=st["InfoGrise"])

        print(f"  ✅ Carte {num} ajoutée")

    doc.save(str(FICHIER_SORTIE))
    print(f"\n✅ Document généré : {FICHIER_SORTIE.name}")
    print(f"   Taille : {FICHIER_SORTIE.stat().st_size / 1_000_000:.1f} Mo")


# =============================================================================
# COMMANDE : statut
# =============================================================================

def cmd_statut(args):
    sauvegarde = charger_sauvegarde()
    cartes     = grouper_images(DOSSIER)

    total    = len(cartes)
    faites   = len(sauvegarde)
    restantes= total - faites
    erreurs  = sum(1 for d in sauvegarde.values() if "erreur" in d)

    print(f"\n📊 État d'avancement")
    print(f"   Total cartes    : {total}")
    print(f"   Analysées       : {faites}  ({faites*100//total if total else 0}%)")
    print(f"   Restantes       : {restantes}")
    print(f"   Erreurs         : {erreurs}")

    if restantes > 0:
        nums_restantes = sorted(
            [n for n in cartes.keys() if str(n) not in sauvegarde]
        )
        print(f"   Cartes restantes : {nums_restantes[:20]}{'...' if len(nums_restantes)>20 else ''}")
        print(f"\n   → python3 generer_synthese_cartes.py analyser")
    else:
        print(f"\n   → python3 generer_synthese_cartes.py generer")

    if erreurs > 0:
        print(f"\n   Cartes en erreur (à réanalyser avec --forcer) :")
        for num_str, d in sauvegarde.items():
            if "erreur" in d:
                print(f"     Carte {num_str} : {d['erreur'][:60]}")


# =============================================================================
# POINT D'ENTRÉE
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Synthèse des cartes postales Lemonnier",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples :
  python3 generer_synthese_cartes.py statut
  python3 generer_synthese_cartes.py analyser --de 1 --a 30
  python3 generer_synthese_cartes.py analyser --cartes 42,43,44
  python3 generer_synthese_cartes.py analyser
  python3 generer_synthese_cartes.py generer
        """
    )
    sub = parser.add_subparsers(dest="commande")

    # analyser
    p_a = sub.add_parser("analyser", help="Analyser les cartes avec l'IA")
    p_a.add_argument("--cartes",  help="Numéros séparés par virgule (ex: 1,5,100)")
    p_a.add_argument("--de",      type=int, help="Numéro de début")
    p_a.add_argument("--a",       type=int, help="Numéro de fin")
    p_a.add_argument("--forcer",  action="store_true", help="Réanalyser même si déjà fait")

    # generer
    sub.add_parser("generer", help="Générer le document Word depuis la sauvegarde")

    # statut
    sub.add_parser("statut", help="Voir l'état d'avancement")

    args = parser.parse_args()

    if args.commande == "analyser":
        cmd_analyser(args)
    elif args.commande == "generer":
        cmd_generer(args)
    elif args.commande == "statut":
        cmd_statut(args)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
